from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ORENDALIS_EXECUTIVE_INTELLIGENCE_EPISODE_1_PRODUCTION_PACKAGE.docx"
MD = ROOT / "ORENDALIS_EXECUTIVE_INTELLIGENCE_EPISODE_1_PRODUCTION_PACKAGE.md"
IMAGE = ROOT.parents[2] / "frontend/public/brand/orendalis-social-preview.png"

NAVY = "0B1220"
BLACK = "070A0F"
INK = "17191C"
WHITE = "F7F9FC"
BLUE = "6D8CFF"
MINT = "7DE2C6"
STONE = "E8DFD3"
MUTED = "68717D"
LIGHT = "F2F4F7"
GOLD = "936B3F"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.82)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.header_distance = sec.footer_distance = Inches(0.35)

def set_font(run, name="Arial", size=10.5, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size, normal.font.color.rgb = "Arial", Pt(10.5), RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
for name, size, color, before, after in [
    ("Title", 30, NAVY, 0, 8), ("Subtitle", 13, MUTED, 0, 12),
    ("Heading 1", 19, NAVY, 18, 9), ("Heading 2", 14, BLUE, 14, 6),
    ("Heading 3", 11.5, NAVY, 10, 4)]:
    st = styles[name]
    st.font.name, st.font.size, st.font.color.rgb = "Arial", Pt(size), RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(before), Pt(after)
    st.paragraph_format.keep_with_next = True
for name in ["List Bullet", "List Number"]:
    st = styles[name]
    st.font.name, st.font.size = "Arial", Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.18

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None: tcMar.append(node)

def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))

def add_p(text="", style=None, bold_lead=None, color=None, size=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, bold=True, color=color or INK, size=size or 10.5)
        r = p.add_run(text[len(bold_lead):]); set_font(r, color=color or INK, size=size or 10.5)
    else:
        r = p.add_run(text); set_font(r, color=color or INK, size=size or 10.5)
    if align is not None: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    return p

def add_bullets(items):
    for item in items: add_p(item, "List Bullet")

def add_callout(label, text, fill=LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    mark_header_row(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False; table.columns[0].width = Inches(6.58)
    cell = table.cell(0,0); shade(cell, fill); set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper()); set_font(r, size=8.5, color=accent, bold=True)
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
    text_color = WHITE if fill in {NAVY, BLACK} else NAVY
    r = p2.add_run(text); set_font(r, size=12, color=text_color, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    mark_header_row(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for i, (head, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]; cell.width = Inches(width); shade(cell, NAVY)
        set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; r = p.add_run(head); set_font(r, size=8.5, color=WHITE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width); set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 0: shade(cells[i], "F7F8FA")
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value)); set_font(r, size=8.4, color=INK, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def page_break(): doc.add_page_break()

def footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ORENDALIS  /  EXECUTIVE INTELLIGENCE  /  EPISODE 1"); set_font(r, size=7.5, color=MUTED, bold=True)

footer(sec)

# Cover
if IMAGE.exists():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    picture = p.add_run().add_picture(str(IMAGE), width=Inches(6.7))
    picture._inline.docPr.set("descr", "Minimal warm-stone architecture opening toward a bright horizon, expressing clarity and forward direction.")
add_p("ORENDALIS", color=BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_p("The Executive Hiring Process Is Broken", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_p("Episode 1 production package", style="Subtitle", align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
add_callout("Creative thesis", "The problem is not a shortage of jobs. The problem is a shortage of intelligence.", fill=NAVY, accent=MINT)
add_p("Master film: 84 seconds  |  Primary channel: LinkedIn  |  Language: English  |  Version: 1.0", color=MUTED, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add_p("Prepared for production  ·  21 July 2026", color=MUTED, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()
add_p("1. Creative concept", "Heading 1")
add_p("The Missing Layer", "Heading 2")
add_p("Executives already have access to jobs, recruiters, networks, documents, calendars, notes and advice. What they lack is a coherent intelligence layer that turns fragmented evidence into a defensible decision. Episode 1 exposes that gap before Orendalis appears.")
add_callout("Audience realization", "I have not been managing a career. I have been managing fragments.")
add_p("The film is a category manifesto, not a product tour. It makes the broken system feel familiar, compresses its hidden cost into one escalating sequence, then introduces Executive Intelligence as the missing alternative.")
add_p("Strategic objective", "Heading 2")
add_bullets([
    "Stop an executive within the first three seconds through recognition, not spectacle.",
    "Replace passive familiarity with discomfort: the old workflow should suddenly feel unacceptable.",
    "Earn the Orendalis reveal only after the viewer understands the category problem.",
    "Leave one durable message: most platforms help people apply; Orendalis helps executives decide.",
])
add_p("Success test", "Heading 2")
add_p("A successful viewer does not merely remember the product. They can repeat the argument in their own words: executive hiring is fragmented, application volume is not intelligence, and important career decisions deserve evidence, context and continuity.")

add_p("2. Film treatment", "Heading 1")
add_p("Recognition", "Heading 2")
add_p("We enter on a black screen with the soft mechanical sound of a cursor. A familiar sentence appears: “Looking for your next executive role?” It is deleted before the viewer can settle. The correction lands: “That’s not the real problem.”")
add_p("Frustration", "Heading 2")
add_p("The screen becomes a controlled accumulation of fragments: job tabs, recruiter messages, CV filenames, saved links, spreadsheet cells, interview notes and compensation research. Nothing is shown as a literal collage for long. Each fragment arrives cleanly, then competes for attention until the system becomes visually impossible.")
add_p("Curiosity", "Heading 2")
add_p("The noise stops. One line remains: “The problem isn’t access.” A second line resolves beneath it: “It’s intelligence.” The soundtrack opens; space returns.")
add_p("Relief", "Heading 2")
add_p("For the first time, the camera moves through a coherent system. Company Intelligence, Opportunity Intelligence, Atlas and the Executive Workspace appear as connected evidence—not as a feature montage. The interface remains readable but never fully explained.")
add_p("Vision", "Heading 2")
add_p("The film leaves the interface and returns to black. The category arrives first, followed by the brand. The final impression is not software. It is a more serious standard for executive decisions.")

page_break()
add_p("3. Master structure and timing", "Heading 1")
timeline = [
    ("00:00–00:03", "Interruption", "Cursor. Familiar question typed, then deleted.", "Stop the scroll."),
    ("00:03–00:09", "Recognition", "The correction replaces it.", "Reframe the problem."),
    ("00:09–00:26", "Fragmentation", "Career tools accumulate with increasing tempo.", "Create productive discomfort."),
    ("00:26–00:36", "Cost", "Tabs freeze; key facts contradict or disappear.", "Make fragmentation consequential."),
    ("00:36–00:45", "Pivot", "Noise collapses into two sentences.", "Name the missing layer."),
    ("00:45–01:07", "Intelligence", "Connected Orendalis evidence and decision views.", "Demonstrate relief without teaching."),
    ("01:07–01:17", "Difference", "Apply versus Decide contrast.", "Own the category distinction."),
    ("01:17–01:24", "End card", "Category, mark, brand and URL.", "Create memory and action."),
]
add_table(["Time", "Movement", "Picture", "Purpose"], timeline, [1.0, 1.15, 2.85, 1.58])
add_p("Pacing law", "Heading 2")
add_p("The first half accelerates; the second half breathes. Fragmentation is cut on rhythm. Intelligence is revealed through slower, deliberate motion. The viewer must feel the difference before it is described.")

add_p("4. Final voiceover", "Heading 1")
add_callout("Performance", "Founder voice. 118–124 words per minute. Calm, observant and certain. Never perform urgency; let the edit create it.", fill="EEF1FF")
vo = [
    "Looking for your next executive role?",
    "That’s not the real problem.",
    "The roles are everywhere.",
    "LinkedIn. Job boards. Recruiters. Company sites.",
    "But the truth about a career decision lives somewhere else.",
    "In a CV version. A saved link. An interview note. A salary file. A conversation you meant to revisit.",
    "More access created more fragments. Not more clarity.",
    "Because the problem isn’t a shortage of jobs.",
    "It’s a shortage of intelligence.",
    "What if every opportunity arrived with context?",
    "Why it fits. Why it may not. What is confirmed. What is unknown. What deserves your attention next.",
    "That is Executive Intelligence.",
    "Most platforms help people apply.",
    "Orendalis helps executives decide.",
    "The executive hiring process is broken.",
    "It’s time to expect better.",
]
for line in vo:
    p = add_p(line, after=5)
    p.paragraph_format.left_indent = Inches(0.28)
add_p("Estimated spoken duration: 71–76 seconds, leaving 8–13 seconds for pauses, silence and the final card.", color=MUTED, size=9)

page_break()
add_p("5. Shot-by-shot storyboard", "Heading 1")
scenes = [
    ("01 · 00:00–00:03", "Black field. A white cursor blinks twice. “Looking for your next executive role?” types in, then is selected.", "VO: Looking for your next executive role?", "Dry keystrokes. No music."),
    ("02 · 00:03–00:09", "The sentence deletes. “That’s not the real problem.” appears one phrase at a time. Hold one full second.", "VO repeats the correction.", "A low sub pulse enters after “problem.”"),
    ("03 · 00:09–00:15", "Three precise browser-tab silhouettes slide into a dark spatial plane: LinkedIn, job board, company careers. No third-party logos needed.", "The roles are everywhere…", "Tempo begins at 74 BPM."),
    ("04 · 00:15–00:26", "Fragments multiply: recruiter email subject, CV_v12_FINAL.pdf, spreadsheet cells, bookmarked role, interview note, calendar hold.", "But the truth… lives somewhere else.", "Each fragment gets a distinct tactile click."),
    ("05 · 00:26–00:36", "Camera pulls back. The fragments form an impossible field. Two salary figures disagree; a role closes; an important note falls outside frame.", "More access created more fragments…", "Pulse tightens; high-frequency texture rises."),
    ("06 · 00:36–00:45", "Everything cuts to black. “NOT A SHORTAGE OF JOBS.” appears, then refines to “A SHORTAGE OF INTELLIGENCE.”", "Because the problem isn’t…", "Complete silence for 12 frames before the second line."),
    ("07 · 00:45–00:52", "A single mint origin point appears. Quiet Ascent path draws upward. A clean Opportunity Intelligence card resolves around it.", "What if every opportunity arrived with context?", "A warm harmonic bed opens."),
    ("08 · 00:52–00:59", "Company Intelligence: evidence, freshness and unknowns enter in three layers. Confirmed / Estimated / Unknown remain visually distinct.", "Why it fits. Why it may not…", "Slow, soft UI ticks."),
    ("09 · 00:59–01:07", "Atlas recommendation panel reveals reasons for, reasons against, confidence and next action. Cursor rests; no chat typing spectacle.", "What deserves your attention next.", "Music reaches emotional clarity, not a climax."),
    ("10 · 01:07–01:13", "Split screen. Left: APPLY, repeated and receding. Right: DECIDE, singular and stable.", "Most platforms help people apply…", "One restrained impact on “decide.”"),
    ("11 · 01:13–01:17", "The Orendalis system recedes into a quiet horizon. “THE EXECUTIVE HIRING PROCESS IS BROKEN.”", "The executive hiring process is broken.", "Music strips back to the opening pulse."),
    ("12 · 01:17–01:24", "Black. Executive Intelligence. Quiet Ascent mark. ORENDALIS. orendalis.com. Fade in that order.", "It’s time to expect better.", "Resolve; leave 18 frames of silence after fade."),
]
for title, picture, words, sound in scenes:
    add_p(title, "Heading 2")
    add_p("PICTURE  "+picture, bold_lead="PICTURE  ")
    add_p("WORDS  "+words, bold_lead="WORDS  ")
    add_p("SOUND  "+sound, bold_lead="SOUND  ", color=MUTED)

page_break()
add_p("6. On-screen text", "Heading 1")
add_p("Use text as argument, not captions. Never display more than nine words in a statement frame. Preserve silence around the strongest lines.")
text_rows = [
    ("00:00", "Looking for your next executive role?", "Typed; deleted"),
    ("00:04", "That’s not the real problem.", "Centered; 1.2 s hold"),
    ("00:37", "NOT A SHORTAGE OF JOBS.", "Condensed tracking; 0.8 s"),
    ("00:40", "A SHORTAGE OF INTELLIGENCE.", "Primary thesis; 2.0 s"),
    ("00:54", "CONFIRMED  /  ESTIMATED  /  UNKNOWN", "Evidence-state microtype"),
    ("01:08", "APPLY", "Left; repeated"),
    ("01:10", "DECIDE", "Right; singular"),
    ("01:14", "THE EXECUTIVE HIRING PROCESS IS BROKEN.", "Final argument"),
    ("01:18", "EXECUTIVE INTELLIGENCE.", "Category reveal"),
    ("01:20", "ORENDALIS", "Brand reveal"),
    ("01:22", "orendalis.com", "Action"),
]
add_table(["Cue", "Text", "Treatment"], text_rows, [0.8, 3.7, 2.08])

add_p("7. Visual identity", "Heading 1")
add_p("Typography", "Heading 2")
add_bullets([
    "Primary: Geist Sans Variable. Headlines use 520–620 weight, tight optical spacing, never all-caps except thesis frames.",
    "Evidence and system labels: Geist Mono, 11–14 px equivalent, generous tracking.",
    "Fallback: Inter Variable and IBM Plex Mono. Do not mix more than two families.",
])
add_p("Color palette", "Heading 2")
palette = [
    ("Obsidian", "#070A0F", "Negative space and manifesto frames"),
    ("Deep ink", "#0B1220", "Interface stage and brand field"),
    ("Signal white", "#F7F9FC", "Primary language"),
    ("Decision blue", "#6D8CFF", "Connections and intelligence"),
    ("Evidence mint", "#7DE2C6", "Confirmed facts and origin point"),
    ("Warm stone", "#E8DFD3", "Human warmth; maximum 10% of frame"),
]
add_table(["Token", "Value", "Use"], palette, [1.3, 1.2, 4.08])
add_p("The Quiet Ascent mark", "Heading 2")
add_p("Use only in Scene 7 as an abstract path and in the final brand reveal. Never open with it. Never animate it as a generic loading spinner. Its motion is a deliberate upward decision path originating from the executive—the mint point—not a symbol of speed.")

add_p("8. Motion and camera language", "Heading 1")
add_bullets([
    "Camera: controlled 2.5D depth, 35–50 mm virtual lens, slow lateral moves, no dramatic orbiting.",
    "Fragment phase: 6–10 frame entrances, hard editorial cuts, shallow depth changes, increasing density.",
    "Intelligence phase: 16–24 frame eased entrances, aligned grids, stable horizon, readable evidence.",
    "Easing: custom cubic 0.22 / 1 / 0.36 / 1. Avoid bounce, elastic motion, parallax excess and cursor chasing.",
    "Transitions: deletion, match cuts, controlled collapse and evidence-line draws. No zoom tunnels or glitch presets.",
    "Texture: 1–2% monochromatic grain; subtle bloom only on blue and mint signal elements.",
])

add_p("9. Music and sound", "Heading 1")
add_p("Music direction", "Heading 2")
add_p("Minimal electronic score at approximately 74 BPM. Start with a dry sub pulse and a faint mechanical texture. Introduce warm, sustained harmony only when intelligence replaces fragmentation. The score should feel designed, not emotionalized.")
add_bullets([
    "Search terms: minimal electronic documentary, restrained pulse, modern design film, warm analog resolution.",
    "Avoid: trailer rises, cinematic booms, inspirational piano, corporate ukulele, heroic strings and vocal samples.",
    "License: use an original composition or a commercial track with worldwide paid-social and perpetual brand-film rights.",
])
add_p("Sound design", "Heading 2")
add_p("Build a small sonic vocabulary: cursor, keystroke, document latch, evidence tick, low pulse and a single decision impact. Every sound must correspond to an action or change in meaning. Silence is an authored element in Scenes 2, 6 and 12.")

page_break()
add_p("10. UI capture direction", "Heading 1")
add_p("Use real Orendalis interface states and truthful data. Capture at 4K or higher, 100% browser zoom, no browser chrome, no personal data and no invented employer claims. Replace sensitive information with purpose-built fictional production records clearly isolated from production data.")
add_p("Required captures", "Heading 2")
add_bullets([
    "Opportunity Intelligence: one role with confirmed evidence, a visible unknown and a meaningful confidence explanation.",
    "Company Intelligence: freshness, provenance and hiring evidence—not fabricated financial or cultural claims.",
    "Atlas: reasons for, reasons against, unknowns, confidence and one recommended next action.",
    "Executive Workspace: continuity from evidence to decision; show Pursue / Watch / Skip only if legible.",
    "Executive Rooms: at most a one-second atmospheric glimpse; never turn the film into a collaboration demo.",
])
add_p("Capture rule", "Heading 2")
add_callout("Truth over spectacle", "If a UI state cannot be shown truthfully, omit it. The film may compress time and space; it may never fabricate evidence or outcomes.", fill="F5EFE7", accent=GOLD)

add_p("11. AI image and video prompts", "Heading 1")
add_p("These prompts create atmospheric connective tissue only. Product UI must be captured from the actual application.")
prompts = [
    ("Fragment field", "Dark editorial 3D space containing restrained floating browser tabs, document edges, calendar blocks and spreadsheet cells; premium product-film lighting; obsidian background; white typography fragments; periwinkle signal lines; subtle grain; no people; no logos; 16:9.", "No holograms, robots, neon cyberpunk, stock-office imagery, readable third-party brands or clutter."),
    ("Collapse to clarity", "Thousands of thin information planes align into one precise horizontal evidence line; black cinematic field; slow deliberate movement; architectural composition; soft periwinkle and mint light; elegant, minimal, high contrast; 16:9.", "No explosions, vortexes, speed ramps, lens flares or sci-fi interfaces."),
    ("Intelligence horizon", "A single mint origin point connects through a restrained ascending path into calm evidence cards on a deep ink horizon; premium editorial motion design; precise grid; subtle depth; quiet confidence; 16:9.", "No charts without labels, no fake data, no glowing brain, no humanoid AI."),
    ("End-card texture", "Near-black field with barely visible warm stone material texture, controlled falloff, central negative space for white typography and logo, premium film grain, no objects, 16:9.", "No gradients that reduce text contrast, no visible symbols or text."),
]
for name, prompt, negative in prompts:
    add_p(name, "Heading 2")
    add_p("PROMPT  "+prompt, bold_lead="PROMPT  ")
    add_p("NEGATIVE  "+negative, bold_lead="NEGATIVE  ", color=MUTED)
add_p("Video generation instruction", "Heading 2")
add_p("Generate 6–8 second plates at 24 fps with locked or extremely slow camera movement. No text generation. No autonomous scene cuts. Preserve negative space for typography. Produce two motion intensities: controlled and almost still. The editor—not the model—determines timing.")

add_p("12. Editorial instructions", "Heading 1")
add_bullets([
    "Master timeline: 3840×2160, 24 fps, 48 kHz audio, 84 seconds.",
    "LinkedIn delivery: 1920×1080 and 1080×1350. Recompose, do not crop mechanically.",
    "Captions: burned-in optional version plus SRT. Sentence case, maximum two lines, WCAG-level contrast.",
    "Brand reveal must not occur before 00:45. Logo must not appear before Scene 7.",
    "UI must remain on screen long enough to understand one idea, not to read every field.",
    "Cut any shot that exists only to prove a feature. Preserve argument, tension and relief.",
    "Export review masters in ProRes 422 HQ; web masters in H.264 high profile, 15–25 Mbps.",
])

add_p("13. Thumbnail concepts", "Heading 1")
thumbs = [
    ("The correction", "Black field. Small deleted line: LOOKING FOR A ROLE? Large white line: THAT’S NOT THE REAL PROBLEM.", "Highest curiosity; recommended."),
    ("The category problem", "Fragmented tabs recede into darkness. Center statement: NOT MORE JOBS. MORE INTELLIGENCE.", "Best for message clarity."),
    ("Apply / Decide", "Vertical split. APPLY repeats in gray on left; DECIDE stands alone in white and mint on right.", "Best for category differentiation."),
]
add_table(["Concept", "Composition", "Use"], thumbs, [1.25, 3.9, 1.43])
add_p("Recommended thumbnail", "Heading 2")
add_p("Use “The correction.” It continues the opening device, creates an unresolved question and does not depend on brand recognition. Keep the logo off the thumbnail unless platform testing proves recognition improves completion rate.")

add_p("14. LinkedIn launch copy", "Heading 1")
add_callout("Opening line", "Executives do not have a job-access problem. They have an intelligence problem.", fill=NAVY, accent=MINT)
post = """For years, we have accepted a fragmented way of managing executive careers.

Jobs in one place. Recruiter conversations in another. CV versions, interview notes, company research, salary files and decisions scattered everywhere else.

More access did not create more clarity.

Because the problem is not a shortage of jobs. It is a shortage of intelligence.

Most platforms help people apply.

ORENDALIS helps executives decide.

Today we are introducing a different category: Executive Intelligence.

Watch Episode 1: The Executive Hiring Process Is Broken.

orendalis.com"""
for paragraph in post.split("\n\n"):
    add_p(paragraph)
add_p("Recommended post treatment", "Heading 2")
add_p("Upload the video natively. Use the first sentence before the fold. Do not list features. Do not use more than three hashtags; preferred: #ExecutiveIntelligence #Leadership #Careers. Founder commentary should be the first comment, not appended to the launch copy.")

page_break()
add_p("15. Production plan", "Heading 1")
plan = [
    ("1", "Voice reference", "Record two complete founder reads and one whisper-quiet read.", "Clean 48 kHz WAV"),
    ("2", "Animatic", "Build timing with typography, temp UI and temp score.", "84-second approval cut"),
    ("3", "UI capture", "Capture only approved truthful states after animatic lock.", "4K lossless plates"),
    ("4", "Motion design", "Create fragmentation and intelligence systems as reusable comps.", "Picture-lock candidate"),
    ("5", "Sound", "Original score, action-based sound design and final founder VO.", "Stereo mix + stems"),
    ("6", "Finishing", "Grade, captions, legal/source review, mobile safe-zone review.", "Master + social exports"),
    ("7", "Validation", "Test first 3 seconds, thesis recall and completion with executives.", "Go/no-go evidence"),
]
add_table(["Stage", "Work", "Acceptance", "Output"], plan, [0.55, 1.25, 3.25, 1.53])
add_p("Voice recording notes", "Heading 2")
add_bullets([
    "Record in a quiet, non-reflective room; microphone 15–20 cm from speaker, slightly off axis.",
    "Do not imitate a commercial narrator. Speak to one respected executive across a table.",
    "Underline only five words in performance: problem, fragments, intelligence, apply, decide.",
    "Keep breaths where they carry thought. Remove mouth noise, not humanity.",
])
add_p("Legal and trust review", "Heading 2")
add_bullets([
    "Use no third-party trademarks or platform screenshots without permission; abstract the fragmented sources.",
    "Use only licensed music, typefaces and generated assets with documented commercial rights.",
    "Do not imply employment outcomes, employer partnerships, application automation or guaranteed recommendations.",
    "Validate every product claim against the live application before picture lock.",
])

page_break()
add_p("16. Acceptance standard", "Heading 1")
acceptance = [
    ("Three-second stop", "Viewer can state the interrupted question or correction after one viewing."),
    ("Problem recognition", "Viewer describes fragmentation before mentioning features."),
    ("Category clarity", "Viewer understands Executive Intelligence as evidence-backed decision support."),
    ("Trust", "No unsupported fact, fabricated UI evidence or exaggerated outcome appears."),
    ("Emotional arc", "The second half feels calmer and more coherent than the first."),
    ("Brand timing", "Orendalis reveal feels earned and occurs only after the thesis."),
    ("Mobile legibility", "All critical text remains readable at 390 px width and in 1080×1350 delivery."),
    ("Message recall", "Primary recall: “The executive hiring process is broken.”"),
]
add_table(["Gate", "Pass condition"], acceptance, [1.75, 4.83])
add_callout("Final creative test", "If removing the Orendalis logo still leaves a film executives want to share, the message is strong enough. Then—and only then—does the brand earn the ending.", fill="EEF1FF")

add_p("17. Episode 1 identity rules", "Heading 1")
add_p("These rules exist only to keep Episode 1 coherent and reusable. They are not production work for later episodes.")
add_bullets([
    "Begin with a human truth, never a product claim.",
    "Use black as thinking space, not as a fashionable dark theme.",
    "Let tension accelerate and intelligence decelerate.",
    "Use blue for connection, mint for confirmed evidence and white for argument.",
    "Keep the founder voice calm, specific and free of startup language.",
    "Reveal the category before the brand; reveal the brand before the URL.",
    "End with a statement worth remembering, not a list of features.",
])
add_p("No future episode has been scripted, storyboarded or produced in this package.", color=MUTED, size=9, after=18)
add_callout("Episode 1 final line", "THE EXECUTIVE HIRING PROCESS IS BROKEN. IT’S TIME TO EXPECT BETTER.", fill=NAVY, accent=MINT)

# Running header
for section in doc.sections:
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("THE MISSING LAYER  ·  EPISODE 1 PRODUCTION PACKAGE")
    set_font(r, size=7.5, color=MUTED, bold=True)

doc.core_properties.title = "ORENDALIS — The Executive Hiring Process Is Broken"
doc.core_properties.subject = "Episode 1 production package for the Executive Intelligence launch film"
doc.core_properties.author = "Orendalis"
doc.core_properties.keywords = "Orendalis, Executive Intelligence, launch film, production package"
doc.save(OUT)

# A concise Markdown source alongside the production artifact.
MD.write_text("""# ORENDALIS — Episode 1 Production Package

## The Executive Hiring Process Is Broken

**Creative concept:** The Missing Layer  
**Master duration:** 84 seconds  
**Thesis:** The problem is not a shortage of jobs. The problem is a shortage of intelligence.

The complete, visually designed production package is maintained in:

`ORENDALIS_EXECUTIVE_INTELLIGENCE_EPISODE_1_PRODUCTION_PACKAGE.docx`

It contains the approved creative concept, treatment, timing, final voiceover, shot-by-shot storyboard, on-screen copy, visual and motion system, music and sound direction, UI capture plan, AI prompts, thumbnail concepts, LinkedIn launch copy, production workflow, legal/trust gates, and acceptance standard.

No future episode is designed or produced in this package.
""", encoding="utf-8")
print(OUT)
